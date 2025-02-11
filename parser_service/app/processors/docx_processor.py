"""DOCX document processor implementation"""

from typing import Dict, Any, Optional
import logging
from datetime import datetime
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from .base import BaseDocumentProcessor

logger = logging.getLogger(__name__)


class DOCXProcessor(BaseDocumentProcessor):
    """Processor for DOCX documents"""

    def _extract_text(self, file_path: str) -> str:
        """Extract text content from DOCX document"""
        try:
            doc = Document(file_path)
            paragraphs = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            paragraphs.append(text)

            return "\n\n".join(paragraphs)
        except PackageNotFoundError:
            logger.error(f"Invalid or corrupted DOCX file: {file_path}")
            raise
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX {file_path}: {str(e)}")
            raise

    def _extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from DOCX document"""
        try:
            doc = Document(file_path)
            core_props = doc.core_properties
            # Count pages (approximate based on sections)
            page_count = len(doc.sections)

            # Count words (approximate based on runs)
            word_count = sum(
                len(paragraph.text.split()) for paragraph in doc.paragraphs
            )

            # Get metadata from core properties
            return {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "created_at": self._convert_docx_date(core_props.created),
                "modified_at": self._convert_docx_date(core_props.modified),
                "last_modified_by": core_props.last_modified_by or "",
                "page_count": page_count,
                "word_count": word_count,
                "category": core_props.category or "",
                "comments": core_props.comments or "",
                "identifier": core_props.identifier or "",
                "keywords": core_props.keywords or "",
                "language": core_props.language or "",
                "subject": core_props.subject or "",
                "version": core_props.version or "",
                "format": "DOCX",
                "file_path": file_path,
            }
        except PackageNotFoundError:
            logger.error(f"Invalid or corrupted DOCX file: {file_path}")
            raise
        except Exception as e:
            logger.error(f"Failed to extract metadata from DOCX {file_path}: {str(e)}")
            raise

    def _convert_docx_date(self, date_value: Any) -> Optional[datetime]:
        """Convert DOCX date value to datetime"""
        if not date_value:
            return None

        try:
            if isinstance(date_value, datetime):
                return date_value
            if isinstance(date_value, str):
                return datetime.fromisoformat(date_value.replace("Z", "+00:00"))
            return None
        except (ValueError, TypeError) as e:
            logger.warning(f"Failed to parse DOCX date {date_value}: {str(e)}")
            return None
