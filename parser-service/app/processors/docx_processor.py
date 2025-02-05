from typing import Dict, Any
from docx import Document
import io
from datetime import datetime
from .base import BaseDocumentProcessor


class DOCXProcessor(BaseDocumentProcessor):
    """Processor for DOCX documents"""

    def _extract_text(self, content: bytes) -> str:
        """Extract text from DOCX document"""
        doc = Document(io.BytesIO(content))
        extracted_text = ""

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                extracted_text += paragraph.text + "\n"

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    extracted_text += " | ".join(row_text) + "\n"

        return extracted_text

    def _extract_metadata(self, content: bytes) -> Dict[str, Any]:
        """Extract metadata from DOCX document"""
        metadata = {}

        try:
            doc = Document(io.BytesIO(content))
            core_properties = doc.core_properties

            # Extract title
            metadata["title"] = core_properties.title or ""

            # Extract author
            metadata["author"] = core_properties.author or ""

            # Extract dates
            if core_properties.created:
                metadata["created_date"] = core_properties.created
            if core_properties.modified:
                metadata["modified_date"] = core_properties.modified

            # Count pages (approximate by counting sections)
            metadata["page_count"] = len(doc.sections)

            # Additional metadata
            metadata["revision"] = core_properties.revision
            metadata["category"] = core_properties.category or ""
            metadata["comments"] = core_properties.comments or ""
            metadata["keywords"] = (
                [k.strip() for k in core_properties.keywords.split(",")]
                if core_properties.keywords
                else []
            )

        except Exception as e:
            # Log error but continue processing
            print(f"Error extracting DOCX metadata: {str(e)}")

        return metadata
